import requests
import sys
import json
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from concurrent.futures import ThreadPoolExecutor, as_completed
import os
from datetime import datetime
import subprocess

def execute_interaction(interaction_id, token, data):
    """
    Executes the interaction using the provided interaction_id and data object.
    """
    url = f"https://studio-server-production.api.becomposable.com/api/v1/interactions/{interaction_id}/execute"
    headers = {
        "Authorization": f"Bearer {token}",
        "Accept": "application/json"
    }
    
    payload = {
        "data": data,
        "stream": False
    }
    
    try:
        response = requests.post(url, headers=headers, json=payload, timeout=120)
        
        if response.status_code == 200:
            try:
                return response.json()  # Ensure response is parsed as JSON
            except ValueError:
                return {"error": "Invalid JSON response"}
        else:
            return {"error": f"Error: {response.status_code} - {response.text}"}
    except requests.exceptions.Timeout:
        return {"error": "The request timed out. Try increasing the timeout duration."}
    except requests.exceptions.RequestException as e:
        return {"error": f"An error occurred: {e}"}

def read_file_as_json(filename):
    """
    Reads the contents of the given file and formats it as a JSON object.
    """
    try:
        with open(filename, 'r', encoding='utf-8') as file:
            content = file.read()
        return {"transcript": content}
    except FileNotFoundError:
        print(f"Error: The file '{filename}' was not found.")
        sys.exit(1)

def call_second_interaction_parallel(second_interaction_id, token, transcript, first_level_outline):
    """
    Calls the second interaction in parallel for each item in first_level_outline and aggregates the results in order.
    """
    sections = []

    def call_second_outline(outline_item):
        data = {
            "transcript": transcript,
            "first_level_outline": first_level_outline,
            "current_outline": outline_item
        }
        
        result = execute_interaction(second_interaction_id, token, data)
        return {"title": outline_item, "subsections": result.get("result", {}).get("second_level_outline", [])}

    # Parallelize calls with ThreadPoolExecutor
    with ThreadPoolExecutor() as executor:
        future_to_outline = {executor.submit(call_second_outline, item): item for item in first_level_outline}
        
        for future in as_completed(future_to_outline):
            outline_item = future_to_outline[future]
            try:
                section = future.result()
                sections.append(section)
            except Exception as exc:
                print(f"Error occurred while processing outline {outline_item}: {exc}")
    
    # Ensure sections are in the original order
    ordered_sections = sorted(sections, key=lambda x: first_level_outline.index(x["title"]))
    return ordered_sections

def create_word_document(transcript, sections, base_file_name):
    """
    Create a Word document with the given sections and transcript, and add a table of contents.
    """
    doc = Document()

    # Add a message for updating the TOC
    doc.add_paragraph("Table of Contents (Please update the Table of Contents in Word: Right-click -> Update Field)").alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    toc = doc.add_paragraph()
    fldChar = OxmlElement('w:fldChar')
    fldChar.set(qn('w:fldCharType'), 'begin')

    instrText = OxmlElement('w:instrText')
    instrText.text = 'TOC \\o "1-3" \\h \\z \\u'

    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'separate')

    fldChar3 = OxmlElement('w:fldChar')
    fldChar3.set(qn('w:fldCharType'), 'end')

    toc._element.append(fldChar)
    toc._element.append(instrText)
    toc._element.append(fldChar2)
    toc._element.append(fldChar3)

    # Add sections, subsections, and points
    for section in sections:
        # Section title (Chapter)
        doc.add_heading(section['title'], level=1)
        
        # Subsections
        for subsection in section['subsections']:
            doc.add_heading(subsection['title'], level=2)
            
            # Points in bullet list
            for content_item in subsection['content']:
                doc.add_paragraph(content_item['point'], style='ListBullet')

    # Save the document with a unique file name
    word_file_name = f"{base_file_name}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
    doc.save(word_file_name)
    print(f"Document '{word_file_name}' created successfully.")

    return word_file_name

def convert_to_pdf(word_file_name):
    """
    Convert the Word document to a PDF.
    """
    pdf_file_name = word_file_name.replace(".docx", ".pdf")
    
    # Use LibreOffice or Microsoft Word for conversion (example with LibreOffice)
    try:
        subprocess.run(['/Applications/LibreOffice.app/Contents/MacOS/soffice', '--headless', '--convert-to', 'pdf', word_file_name], check=True)
        print(f"PDF '{pdf_file_name}' created successfully.")
    except Exception as e:
        print(f"Error during PDF conversion: {e}")
    
    return pdf_file_name

if __name__ == "__main__":
    if len(sys.argv) != 5:
        print("Usage: python script_name.py <interaction_id> <token> <file_name> <second_interaction_id>")
        sys.exit(1)
    
    interaction_id = sys.argv[1]
    token = sys.argv[2]
    file_name = sys.argv[3]
    second_interaction_id = sys.argv[4]

    # Extract base file name from the transcript file (without extension)
    base_file_name = os.path.splitext(os.path.basename(file_name))[0]
    
    # Step 1: Read the transcript from the file
    data = read_file_as_json(file_name)

    # Step 2: Call the first interaction to get the first-level outline
    result = execute_interaction(interaction_id, token, data)
    
    # Safely check for errors and process only if the response is valid
    if "error" in result:
        print(result["error"])
        sys.exit(1)
    
    first_level_outline = result.get("result", {}).get("first_level_outline", [])

    if not first_level_outline:
        print("No first level outline found.")
    else:
        # Step 3: Call the second interaction in parallel for each outline item
        transcript = data["transcript"]
        sections = call_second_interaction_parallel(second_interaction_id, token, transcript, first_level_outline)
        
        # Step 4: Create Word document with unique name
        word_file_name = create_word_document(transcript, sections, base_file_name)
        
        # Step 5: Convert Word document to PDF

        
        convert_to_pdf(word_file_name)

